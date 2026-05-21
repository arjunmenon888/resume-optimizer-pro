import re
import uuid
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from config import config

# ---------------------------------------------------------------------------
# Section-header detection
# ---------------------------------------------------------------------------
_SECTIONS = [
    ("summary",    re.compile(r"^(professional\s+)?summary|objective|profile|about\s*me?$", re.I)),
    ("experience", re.compile(r"^(professional\s+|work\s+)?experience|employment(\s+history)?|work\s+history$", re.I)),
    ("education",  re.compile(r"^(academic\s+)?education|academic\s+background|qualifications?$", re.I)),
    ("skills",     re.compile(r"^(technical\s+|core\s+|key\s+)?skills|technologies|tools|competencies|expertise$", re.I)),
    ("certifications", re.compile(r"^certifications?|certificates?|licenses?$", re.I)),
]

# Placeholder lines like [Your Education Information]
_PLACEHOLDER_RE = re.compile(r"^\[.+\]$")


def _classify_section(line: str) -> str | None:
    """Return section key if line is a section heading, else None."""
    clean = re.sub(r"[:\-*#_]+$", "", line).strip()
    clean = re.sub(r"^[:\-*#_]+", "", clean).strip()
    for key, pattern in _SECTIONS:
        if pattern.fullmatch(clean):
            return key
    return None


def _clean(lines: list[str]) -> list[str]:
    out = []
    for l in lines:
        s = l.strip()
        if s and not _PLACEHOLDER_RE.match(s):
            out.append(s)
    return out


# ---------------------------------------------------------------------------
# Section-specific parsers
# ---------------------------------------------------------------------------

def _parse_summary(lines: list[str]) -> str:
    return " ".join(_clean(lines))


def _parse_experience(lines: list[str]) -> list[dict]:
    entries: list[dict] = []
    current: dict | None = None
    desc: list[str] = []

    def flush():
        nonlocal current, desc
        if current is not None:
            current["description"] = "\n".join(desc)
            entries.append(current)
            current = None
            desc = []

    for line in _clean(lines):
        is_bullet = line.startswith(("-", "•", "*", "·"))

        if is_bullet:
            if current is not None:
                desc.append(line.lstrip("-•*· ").strip())
            continue

        # "Title | Company | Location | Dates" — all on one line
        if "|" in line:
            parts = [p.strip() for p in line.split("|")]
            # Decide: is first part a title or is this a company line for the current entry?
            first_has_date = bool(re.search(r"\d{4}|present|current|now", parts[0], re.I))
            if current and not current.get("company") and not first_has_date:
                # This is the "Company | Location | Dates" follow-up line
                current["company"] = parts[0]
                for part in reversed(parts[1:]):
                    if re.search(r"\d{4}|present|current|now", part, re.I):
                        current["duration"] = part
                        break
            else:
                # New entry with everything on one line
                flush()
                title = parts[0]
                company = parts[1] if len(parts) > 1 else ""
                duration = ""
                for part in reversed(parts[1:]):
                    if re.search(r"\d{4}|present|current|now", part, re.I):
                        duration = part
                        break
                current = {"title": title, "company": company, "duration": duration, "description": ""}
            continue

        # Non-bullet, no pipe → new job title
        flush()
        current = {"title": line, "company": "", "duration": "", "description": ""}

    flush()
    return entries


def _parse_education(lines: list[str]) -> list[dict]:
    entries: list[dict] = []
    current: dict | None = None

    def flush():
        nonlocal current
        if current is not None:
            entries.append(current)
            current = None

    for line in _clean(lines):
        if "|" in line:
            parts = [p.strip() for p in line.split("|")]
            if current and not current.get("school"):
                current["school"] = parts[0]
                year_match = re.search(r"\b(19|20)\d{2}\b", line)
                if year_match:
                    current["year"] = year_match.group()
            else:
                flush()
                degree = parts[0]
                school = parts[1] if len(parts) > 1 else ""
                year_match = re.search(r"\b(19|20)\d{2}\b", line)
                year = year_match.group() if year_match else ""
                entries.append({"degree": degree, "school": school, "year": year})
        elif re.search(r"\b(bachelor|master|phd|doctor|associate|diploma|degree|b\.s|m\.s|b\.a|m\.a)\b", line, re.I):
            flush()
            year_match = re.search(r"\b(19|20)\d{2}\b", line)
            year = year_match.group() if year_match else ""
            line_clean = re.sub(r"\b(19|20)\d{2}\b", "", line).strip(" |,")
            current = {"degree": line_clean, "school": "", "year": year}
        else:
            flush()
            current = {"degree": line, "school": "", "year": ""}

    flush()
    return entries


def _parse_skills(lines: list[str]) -> list[str]:
    raw: list[str] = []
    for line in _clean(lines):
        # Strip leading bullet/dash
        text = line.lstrip("-•*· ").strip()
        # Split on commas, semicolons, bullets within the line
        parts = re.split(r"[,;•·]", text)
        for p in parts:
            s = p.strip().strip("()")
            if s and len(s) < 60:
                raw.append(s)
    # Deduplicate while preserving order
    seen: set[str] = set()
    out: list[str] = []
    for s in raw:
        key = s.lower()
        if key not in seen:
            seen.add(key)
            out.append(s)
    return out


# ---------------------------------------------------------------------------
# Main parser
# ---------------------------------------------------------------------------

class DocumentService:

    def parse_resume_text(self, text: str) -> dict:
        """
        Parse raw LLM-generated resume text into structured sections.
        Handles section headers like 'Summary:', 'Professional Experience:',
        'Education:', 'Skills:', and placeholder lines like [Your Contact Info].
        """
        lines = [l.rstrip() for l in text.splitlines()]

        # --- Extract name: first non-empty, non-placeholder, non-header line ---
        name = "Your Name"
        for line in lines:
            s = line.strip()
            if s and not _PLACEHOLDER_RE.match(s) and not _classify_section(s):
                name = s.lstrip("#*_ ").rstrip("#*_ :")
                break

        # --- Split lines into named sections ---
        bucket: str = "_header"
        buckets: dict[str, list[str]] = {bucket: []}

        for line in lines:
            s = line.strip().lstrip("#*_ ").rstrip("#*_ :")
            key = _classify_section(s) if s else None
            if key:
                bucket = key
                buckets.setdefault(bucket, [])
            else:
                buckets.setdefault(bucket, []).append(line)

        # Certifications → append to skills
        cert_lines = buckets.get("certifications", [])
        buckets.setdefault("skills", []).extend(cert_lines)

        return {
            "name": name,
            "email": "",
            "phone": "",
            "location": "",
            "summary": _parse_summary(buckets.get("summary", [])),
            "experience": _parse_experience(buckets.get("experience", [])),
            "education": _parse_education(buckets.get("education", [])),
            "skills": _parse_skills(buckets.get("skills", [])),
        }

    # -----------------------------------------------------------------------
    # Word document generation
    # -----------------------------------------------------------------------

    def generate_ats_resume(self, resume_data: dict) -> Document:
        """Generate ATS-optimized Word document from structured resume data."""
        doc = Document()

        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        # Name header
        name_para = doc.add_paragraph(resume_data.get("name", "Your Name"))
        name_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = name_para.runs[0]
        run.font.size = Pt(14)
        run.font.bold = True

        # Contact line — skip blank fields
        contact_parts = [
            v for v in [
                resume_data.get("email", ""),
                resume_data.get("phone", ""),
                resume_data.get("location", ""),
            ] if v
        ]
        if contact_parts:
            contact_para = doc.add_paragraph(" | ".join(contact_parts))
            contact_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph("").space_after = Pt(6)

        # Professional Summary
        if resume_data.get("summary"):
            doc.add_heading("PROFESSIONAL SUMMARY", level=2)
            doc.add_paragraph(resume_data["summary"]).space_after = Pt(10)

        # Professional Experience
        if resume_data.get("experience"):
            doc.add_heading("PROFESSIONAL EXPERIENCE", level=2)
            for exp in resume_data["experience"]:
                if not exp.get("title"):
                    continue
                title_para = doc.add_paragraph(exp["title"])
                if title_para.runs:
                    title_para.runs[0].font.bold = True
                title_para.space_after = Pt(0)

                company = exp.get("company", "")
                duration = exp.get("duration", "")
                company_line = " | ".join(p for p in [company, duration] if p)
                if company_line:
                    cp = doc.add_paragraph(company_line)
                    if cp.runs:
                        cp.runs[0].italic = True
                    cp.space_after = Pt(4)

                description = exp.get("description", "")
                if description:
                    for bullet in description.splitlines():
                        b = bullet.strip()
                        if b:
                            p = doc.add_paragraph(style="List Bullet")
                            p.add_run(b)
                doc.add_paragraph("").space_after = Pt(8)

        # Education
        if resume_data.get("education"):
            doc.add_heading("EDUCATION", level=2)
            for edu in resume_data["education"]:
                if not edu.get("degree"):
                    continue
                dp = doc.add_paragraph(edu["degree"])
                if dp.runs:
                    dp.runs[0].font.bold = True
                dp.space_after = Pt(0)

                school = edu.get("school", "")
                year = edu.get("year", "")
                school_line = " | ".join(p for p in [school, year] if p)
                if school_line:
                    doc.add_paragraph(school_line).space_after = Pt(8)

        # Skills
        if resume_data.get("skills"):
            doc.add_heading("SKILLS", level=2)
            doc.add_paragraph(" • ".join(resume_data["skills"]))

        return doc

    async def save_resume(self, doc: Document, file_name: str = None) -> dict:
        """Save Word document and return metadata."""
        os.makedirs(config.UPLOAD_DIR, exist_ok=True)
        resume_id = str(uuid.uuid4())
        file_name = file_name or f"{resume_id}.docx"
        file_path = os.path.join(config.UPLOAD_DIR, file_name)
        doc.save(file_path)
        return {
            "id": resume_id,
            "file_path": file_path,
            "file_name": file_name,
            "download_url": f"/api/download/{resume_id}",
        }


document_service = DocumentService()
